from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.services.document_loader import SUPPORTED_SUFFIXES, DocumentLoader


@pytest.fixture
def loader() -> DocumentLoader:
    return DocumentLoader()


def test_supported_suffixes() -> None:
    assert {".pdf", ".docx", ".txt"} == SUPPORTED_SUFFIXES


def test_load_txt(loader: DocumentLoader, tmp_path: Path) -> None:
    file_path = tmp_path / "contract.txt"
    file_path.write_text("Contract between Alpha LLC and Beta Ltd.", encoding="utf-8")

    document = loader.load(file_path)

    assert document.format == "txt"
    assert "Alpha LLC" in document.text
    assert document.filename == "contract.txt"


def test_load_docx(loader: DocumentLoader, tmp_path: Path) -> None:
    file_path = tmp_path / "memo.docx"
    doc = Document()
    doc.add_paragraph("Payment deadline: 2026-12-31")
    doc.add_paragraph("Amount: 1,500,000 RUB")
    doc.save(file_path)

    document = loader.load(file_path)

    assert document.format == "docx"
    assert "Payment deadline" in document.text
    assert "1,500,000 RUB" in document.text


def test_load_pdf(loader: DocumentLoader, tmp_path: Path) -> None:
    file_path = tmp_path / "report.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with file_path.open("wb") as pdf_file:
        writer.write(pdf_file)

    with pytest.raises(ValueError, match="No extractable text"):
        loader.load(file_path)


def test_load_missing_file(loader: DocumentLoader, tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError, match="File not found"):
        loader.load(tmp_path / "missing.txt")


def test_load_unsupported_format(loader: DocumentLoader, tmp_path: Path) -> None:
    file_path = tmp_path / "data.csv"
    file_path.write_text("a,b,c", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        loader.load(file_path)


def test_load_empty_txt(loader: DocumentLoader, tmp_path: Path) -> None:
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   \n  ", encoding="utf-8")

    with pytest.raises(ValueError, match="No extractable text"):
        loader.load(file_path)
