from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}


@dataclass(frozen=True, slots=True)
class LoadedDocument:
    path: Path
    filename: str
    format: str
    text: str


class DocumentLoader:
    """Load supported document formats and extract plain text."""

    def load(self, file_path: Path) -> LoadedDocument:
        path = file_path.expanduser().resolve()
        if not path.is_file():
            msg = f"File not found: {path}"
            raise FileNotFoundError(msg)

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
            msg = f"Unsupported file type '{suffix}'. Supported: {supported}"
            raise ValueError(msg)

        text = self._extract_text(path, suffix).strip()
        if not text:
            msg = f"No extractable text found in: {path.name}"
            raise ValueError(msg)

        return LoadedDocument(
            path=path,
            filename=path.name,
            format=suffix.removeprefix("."),
            text=text,
        )

    def _extract_text(self, path: Path, suffix: str) -> str:
        if suffix == ".txt":
            return path.read_text(encoding="utf-8")
        if suffix == ".docx":
            return self._extract_docx(path)
        return self._extract_pdf(path)

    def _extract_docx(self, path: Path) -> str:
        document = Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        return "\n".join(paragraph for paragraph in paragraphs if paragraph)

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(page.strip() for page in pages if page.strip())
